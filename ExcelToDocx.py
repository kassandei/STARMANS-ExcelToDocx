import pandas as pd
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
import os
import subprocess
import sys

def format_number(value, decimals):
    if isinstance(value, (int, float)):
        return f"{value:.{decimals}f}"
    return value

def process_file(input_file, output_file):
    try:
        # Extract the base name of the file
        base_name = os.path.basename(input_file)
        file_name = os.path.splitext(base_name)[0]

        # Check if the file exists
        if not os.path.isfile(input_file):
            raise ValueError("Vybraný soubor neexistuje")

        # Check if the file has a valid Excel extension
        if not (input_file.endswith('.xlsx') or input_file.endswith('.xls')):
            raise ValueError("Vybraný soubor není soubor Excel")

        # Check available sheets
        xls = pd.ExcelFile(input_file, engine='openpyxl')
        sheet_name = xls.sheet_names[0]  # Using the first sheet name

        # Define the ranges to check
        ranges = [
            ('Q:W', 23, 4),  # range Q25:W28
            ('G:M', 27, 4)   # range G29:M32
        ]

        df_list = []
        for col_range, skip_rows, nrows in ranges:
            try:
                temp_df = pd.read_excel(input_file, sheet_name=sheet_name, usecols=col_range, skiprows=skip_rows, nrows=nrows, engine='openpyxl')
                if not temp_df.empty:
                    df_list.append(temp_df)
            except Exception as e:
                print(f"Failed to read range {col_range} with error: {e}")
        
        if not df_list:
            raise ValueError("Nenalezeny žádné platné datové rozsahy v Excel souboru")
        
        # Combine valid dataframes if necessary (this example assumes using the first valid dataframe)
        df = df_list[0]

        # Format the numbers to match significant figures
        decimals = [2, 2, 3, 2, 2, 2, 2]  # Number of decimal places for each column
        for col, dec in zip(df.columns, decimals):
            df[col] = df[col].apply(lambda x: format_number(x, dec))

        # Debugging: Print the dataframe
        print(df)

        # Load the provided Word document
        doc = Document("DOCXFILE.docx")

        # Find the table to replace its content with the new data
        table = doc.tables[0]  # Assuming the table to replace is the first table

        # Get the last four rows in the table
        rows_to_replace = table.rows[-4:]

        # Replace the data in the last four rows
        for row_idx, row_data in enumerate(df.values):
            row_cells = rows_to_replace[row_idx].cells
            for col_idx, cell_value in enumerate(row_data):
                p = row_cells[col_idx].paragraphs[0]
                p.clear()  # Clear existing content in the cell
                run = p.add_run(str(cell_value))
                run.font.name = 'Calibri'
                run.font.size = Pt(11)
                p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Save the modified document
        doc.save(output_file)

        # Automatically open the document
        if sys.platform.startswith('linux'):
            subprocess.Popen(['xdg-open', output_file])
        elif sys.platform.startswith('darwin'):  # macOS
            subprocess.Popen(['open', output_file])
        elif sys.platform.startswith('win'):  # Windows
            os.startfile(output_file)

        print(f'Tabulka uložena do: {output_file}')
    except ValueError as ve:
        print(ve)
    except Exception as e:
        print(f"Nastala chyba při zpracování souboru: {e}")

if __name__ == "__main__":
    if len(sys.argv) != 3:
        print("Usage: python script_name.py <input_file.xlsx> <output_file.docx>")
        sys.exit(1)

    input_file = sys.argv[1]
    output_file = sys.argv[2]
    process_file(input_file, output_file)
